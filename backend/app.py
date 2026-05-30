import os
import re
import json
import fitz
from io import BytesIO
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document

load_dotenv()

app = Flask(__name__)
CORS(app)

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def extract_text_from_pdf(file):
    text = ""
    pdf = fitz.open(stream=file.read(), filetype="pdf")

    for page in pdf:
        text += page.get_text() + "\n"

    return text.strip()


def extract_contact_info(resume_text):
    lines = [line.strip() for line in resume_text.split("\n") if line.strip()]

    email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", resume_text)
    phone_match = re.search(r"(\+?\d[\d\s\-\(\)]{8,}\d)", resume_text)
    linkedin_match = re.search(r"(https?://)?(www\.)?linkedin\.com/[^\s]+", resume_text, re.I)

    possible_name = "Candidate"
    for line in lines[:8]:
        if (
            len(line.split()) >= 2
            and len(line.split()) <= 4
            and "@" not in line
            and not any(char.isdigit() for char in line)
            and "linkedin" not in line.lower()
            and "resume" not in line.lower()
        ):
            possible_name = line
            break

    return {
        "name": possible_name,
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0) if phone_match else "",
        "linkedin": linkedin_match.group(0) if linkedin_match else "",
    }


@app.route("/")
def home():
    return jsonify({"message": "AI Resume Checker Backend Running"})


@app.route("/analyze", methods=["POST"])
def analyze_resume():
    try:
        if "resume" not in request.files:
            return jsonify({"error": "No resume uploaded"}), 400

        resume_file = request.files["resume"]
        job_description = request.form.get("job_description", "")

        resume_text = extract_text_from_pdf(resume_file)

        prompt = f"""
You are an expert resume reviewer, ATS analyst, recruiter, and career coach.

Return ONLY valid JSON. No markdown. No backticks.

Use this structure:

{{
  "match_score": 0,
  "summary": "Short summary.",
  "strengths": ["strength 1", "strength 2", "strength 3"],
  "weaknesses": ["weakness 1", "weakness 2", "weakness 3"],
  "missing_keywords": ["keyword 1", "keyword 2", "keyword 3"],
  "salary_estimate": {{
    "low": "$0",
    "high": "$0",
    "reasoning": "Short explanation."
  }},
  "interview_questions": ["question 1", "question 2", "question 3", "question 4", "question 5"],
  "recommended_resume_changes": ["change 1", "change 2", "change 3"]
}}

Resume:
{resume_text}

Job Description:
{job_description}
"""

        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
        )

        raw_output = response.choices[0].message.content

        try:
            parsed_output = json.loads(raw_output)
        except json.JSONDecodeError:
            parsed_output = {
                "match_score": 0,
                "summary": raw_output,
                "strengths": [],
                "weaknesses": [],
                "missing_keywords": [],
                "salary_estimate": {
                    "low": "N/A",
                    "high": "N/A",
                    "reasoning": "Could not parse salary estimate."
                },
                "interview_questions": [],
                "recommended_resume_changes": []
            }
        parsed_output["resume_text"] = resume_text
        return jsonify(parsed_output)

    except Exception as e:
        print("ERROR:", str(e))
        return jsonify({"error": str(e)}), 500


@app.route("/generate_resume", methods=["POST"])
def generate_resume():
    try:
        data = request.json

        analysis = data.get("analysis", {})
        job_description = data.get("job_description", "")
        resume_text = data.get("resume_text", "")

        contact_info = extract_contact_info(resume_text) if resume_text else {
            "name": "Candidate",
            "email": "",
            "phone": "",
            "linkedin": "",
        }

        prompt = f"""
You are an expert resume writer and ATS optimization specialist.

Create a personalized, ATS-friendly resume based on the job description and resume analysis.

Use this candidate contact information if available:
Name: {contact_info["name"]}
Phone: {contact_info["phone"]}
Email: {contact_info["email"]}
LinkedIn: {contact_info["linkedin"]}

Do NOT use placeholders like [Full Name], [Phone Number], [Email Address], or [LinkedIn URL].
Do NOT invent fake companies, fake degrees, fake certifications, or fake years of experience.
Only improve wording, structure, keywords, and positioning based on the candidate's existing background.

Return the resume in clean plain text.

Include:

1. CONTACT INFORMATION
2. PROFESSIONAL SUMMARY
3. CORE SKILLS
4. EXPERIENCE BULLETS
5. KEYWORDS ADDED
6. FINAL ATS NOTES

Resume Analysis:
{json.dumps(analysis, indent=2)}

Job Description:
{job_description}
"""

        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
        )

        return jsonify({"resume": response.choices[0].message.content})

    except Exception as e:
        print("ERROR:", str(e))
        return jsonify({"error": str(e)}), 500


@app.route("/download_resume", methods=["POST"])
def download_resume():
    try:
        data = request.json
        resume_text = data.get("resume", "")

        document = Document()

        for line in resume_text.split("\n"):
            clean_line = line.strip()

            if not clean_line:
                document.add_paragraph("")
            elif clean_line.isupper() and len(clean_line) < 45:
                document.add_heading(clean_line, level=2)
            elif clean_line.startswith("-"):
                document.add_paragraph(clean_line[1:].strip(), style="List Bullet")
            else:
                document.add_paragraph(clean_line)

        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name="Aiko_Optimized_Resume.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        print("ERROR:", str(e))
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(debug=True)